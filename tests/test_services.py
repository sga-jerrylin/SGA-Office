"""
Service 层单元测试。
测试 doc_builder, excel_handler 的纯业务逻辑。
这些函数不直接调用 COS，可以直接测试输入输出。
"""

import pytest
from io import BytesIO
from docx import Document


# =====================================================
#  DOC-01: render_markdown_to_docx
# =====================================================

class TestRenderMarkdownToDocx:

    def test_basic_heading_and_paragraph(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = "# 测试标题\n\n这是一段正文内容。"
        result = render_markdown_to_docx(md)
        assert isinstance(result, BytesIO)
        # 验证是合法的 docx
        doc = Document(result)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("测试标题" in t for t in texts)
        assert any("正文内容" in t for t in texts)

    def test_markdown_with_table(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = "# Report\n\n| Name | Age |\n|---|---|\n| Alice | 30 |\n| Bob | 25 |"
        result = render_markdown_to_docx(md)
        doc = Document(result)
        # 应该包含一个表格
        assert len(doc.tables) >= 1

    def test_markdown_with_bold_and_italic(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = "This has **bold** and *italic* text."
        result = render_markdown_to_docx(md)
        doc = Document(result)
        assert any("bold" in p.text for p in doc.paragraphs)

    def test_escaped_newlines_converted(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = "Line1\\nLine2\\nLine3"
        result = render_markdown_to_docx(md)
        doc = Document(result)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Line1" in full_text
        assert "Line3" in full_text

    def test_empty_result_is_valid_docx(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = "Just plain text"
        result = render_markdown_to_docx(md)
        result.seek(0)
        # 验证 ZIP 签名 (docx 是 zip 格式)
        assert result.read(2) == b"PK"


# =====================================================
#  DOC-02: fill_docx_template
# =====================================================

class TestFillDocxTemplate:

    def _create_template_bytes(self, text_content: str) -> bytes:
        """创建一个包含占位符的简单 docx 模板"""
        doc = Document()
        doc.add_paragraph(text_content)
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_simple_placeholder_replacement(self):
        from app.services.doc_builder import fill_docx_template
        # 注意：实现使用 {{key}} 无空格格式
        template = self._create_template_bytes("Hello {{name}}, welcome to {{company}}.")
        result = fill_docx_template(template, {"name": "Alice", "company": "ACME"})
        doc = Document(result)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Alice" in full_text
        assert "ACME" in full_text

    def test_unmatched_placeholders_remain(self):
        from app.services.doc_builder import fill_docx_template
        # 注意：实现使用 {{key}} 无空格格式
        template = self._create_template_bytes("Hello {{name}} and {{unknown}}.")
        result = fill_docx_template(template, {"name": "Bob"})
        doc = Document(result)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Bob" in full_text
        # unknown 占位符应保持不变
        assert "{{unknown}}" in full_text


# =====================================================
#  EXC-01: create_excel_from_array
# =====================================================

class TestCreateExcelFromArray:

    def test_basic_creation(self):
        from app.services.excel_handler import create_excel_from_array
        result = create_excel_from_array(
            title="Test Report",
            data=[["Name", "Score"], ["Alice", 95], ["Bob", 87]],
        )
        assert isinstance(result, BytesIO)
        # 验证是合法的 xlsx
        result.seek(0)
        assert result.read(2) == b"PK"

    def test_custom_sheet_name(self):
        from app.services.excel_handler import create_excel_from_array
        from openpyxl import load_workbook
        result = create_excel_from_array(
            title="Report",
            data=[["Col1", "Col2"], ["a", "b"]],
            sheet_name="MySheet",
        )
        wb = load_workbook(result)
        assert "MySheet" in wb.sheetnames


# =====================================================
#  EXC-03: generate_complex_excel
# =====================================================

class TestGenerateComplexExcel:

    def test_multi_sheet_creation(self):
        from app.services.excel_handler import generate_complex_excel
        from openpyxl import load_workbook
        result = generate_complex_excel(
            title="Financial Report",
            sheets_def=[
                {
                    "sheet_name": "Revenue",
                    "headers": ["Month", "Amount"],
                    "data": [["Jan", 10000], ["Feb", 12000]],
                },
                {
                    "sheet_name": "Expenses",
                    "headers": ["Month", "Amount"],
                    "data": [["Jan", 8000], ["Feb", 9000]],
                },
            ],
        )
        wb = load_workbook(result)
        assert "Revenue" in wb.sheetnames
        assert "Expenses" in wb.sheetnames

    def test_with_formula(self):
        from app.services.excel_handler import generate_complex_excel
        from openpyxl import load_workbook
        result = generate_complex_excel(
            title="Sum Test",
            sheets_def=[
                {
                    "sheet_name": "Data",
                    "headers": ["A", "B", "Total"],
                    "data": [[1, 2, "=A2+B2"], [3, 4, "=A3+B3"]],
                },
            ],
        )
        wb = load_workbook(result)
        ws = wb["Data"]
        # 公式应该被保留
        # 找到含有 '=A' 的单元格（公式行）
        has_formula = False
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.startswith("="):
                    has_formula = True
                    break
        assert has_formula


# =====================================================
#  PDF helper: _hex_to_rgb
# =====================================================

class TestHexToRgb:

    def test_black(self):
        from app.services.pdf_manipulator import _hex_to_rgb
        r, g, b = _hex_to_rgb("#000000")
        assert (r, g, b) == (0.0, 0.0, 0.0)

    def test_white(self):
        from app.services.pdf_manipulator import _hex_to_rgb
        r, g, b = _hex_to_rgb("#FFFFFF")
        assert abs(r - 1.0) < 0.01
        assert abs(g - 1.0) < 0.01
        assert abs(b - 1.0) < 0.01

    def test_no_hash(self):
        from app.services.pdf_manipulator import _hex_to_rgb
        r, g, b = _hex_to_rgb("808080")
        assert abs(r - 0.502) < 0.01

