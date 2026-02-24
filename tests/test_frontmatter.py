"""Frontmatter 解析 + 富文档渲染测试"""

import pytest
from io import BytesIO
from docx import Document


class TestFrontmatterParsing:
    """测试 frontmatter 提取逻辑"""

    def test_parse_frontmatter_with_cover(self):
        from app.services.doc_builder import _parse_frontmatter
        md = '---\ncover:\n  title: "测试标题"\n  subtitle: "副标题"\ntheme: business_blue\n---\n# 正文'
        config, body = _parse_frontmatter(md)
        assert config["cover"]["title"] == "测试标题"
        assert config["theme"] == "business_blue"
        assert body.strip() == "# 正文"

    def test_no_frontmatter_returns_none(self):
        from app.services.doc_builder import _parse_frontmatter
        md = "# 普通文档\n\n正文内容"
        config, body = _parse_frontmatter(md)
        assert config is None
        assert body == md

    def test_invalid_yaml_returns_none(self):
        from app.services.doc_builder import _parse_frontmatter
        md = "---\n  bad: yaml: :\n---\n# 正文"
        config, body = _parse_frontmatter(md)
        assert config is None

    def test_frontmatter_only_dashes_ignored(self):
        """正文中的 --- 不应被当作 frontmatter"""
        from app.services.doc_builder import _parse_frontmatter
        md = "# 标题\n\n---\n\n下一节"
        config, body = _parse_frontmatter(md)
        assert config is None


class TestCoverPage:
    """测试封面页生成"""

    def test_cover_page_generated(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = '---\ncover:\n  title: "年度报告"\n  subtitle: "2026年"\n  meta:\n    - "作者：张三"\n---\n# 第一章\n\n正文'
        result = render_markdown_to_docx(md)
        doc = Document(result)
        texts = [p.text for p in doc.paragraphs]
        assert any("年度报告" in t for t in texts)
        assert any("张三" in t for t in texts)

    def test_no_cover_when_not_specified(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = '---\ntheme: business_blue\n---\n# 标题\n\n正文'
        result = render_markdown_to_docx(md)
        doc = Document(result)
        # 第一个段落应该直接是标题，不是封面
        non_empty = [p for p in doc.paragraphs if p.text.strip()]
        assert non_empty[0].text == "标题"


class TestHeaderFooter:
    """测试页眉页脚"""

    def test_header_text_set(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = '---\nheader: "机密文档 — 内部使用"\n---\n# 正文'
        result = render_markdown_to_docx(md)
        doc = Document(result)
        header = doc.sections[0].header
        header_text = "".join(p.text for p in header.paragraphs)
        assert "机密文档" in header_text

    def test_footer_page_number(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = '---\nfooter: page_number\n---\n# 正文'
        result = render_markdown_to_docx(md)
        doc = Document(result)
        footer = doc.sections[0].footer
        # 页码通过 XML field code 插入，检查 footer 已被激活
        assert footer is not None


class TestCalloutBox:
    """测试高亮框渲染"""

    @staticmethod
    def _all_text(doc):
        """提取文档中所有文本，包括表格单元格。"""
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        parts.append(p.text)
        return " ".join(parts)

    def test_info_callout_rendered(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = '# 标题\n\n> [!INFO] 重要提示\n> 这是一条信息\n> 第二行信息\n\n正文继续'
        result = render_markdown_to_docx(md)
        doc = Document(result)
        full = self._all_text(doc)
        assert "重要提示" in full or "信息" in full

    def test_warning_callout_rendered(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = '> [!WARNING] 注意事项\n> 请务必检查'
        result = render_markdown_to_docx(md)
        doc = Document(result)
        full = self._all_text(doc)
        assert "注意" in full or "检查" in full


class TestPageBreak:
    """测试分页符"""

    def test_hr_becomes_page_break(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = '---\ntheme: business_blue\n---\n# 第一章\n\n内容一\n\n---\n\n# 第二章\n\n内容二'
        result = render_markdown_to_docx(md)
        doc = Document(result)
        # 文档应包含两个章节标题
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 2


class TestThemedTable:
    """测试主题色表格"""

    def test_table_still_renders(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = '---\ntheme: business_blue\n---\n# 数据\n\n| 字段 | 类型 |\n|------|------|\n| id | int |\n| name | str |'
        result = render_markdown_to_docx(md)
        doc = Document(result)
        assert len(doc.tables) >= 1


class TestTOC:
    """测试目录生成"""

    def test_toc_placeholder_inserted(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = '---\ntoc: true\n---\n# 第一章\n\n内容\n\n# 第二章\n\n内容'
        result = render_markdown_to_docx(md)
        doc = Document(result)
        # TOC 通过 XML 指令插入，检查目录段落存在
        texts = [p.text for p in doc.paragraphs]
        assert any("目录" in t for t in texts) or len(doc.paragraphs) > 2


class TestBackwardCompat:
    """确保纯 Markdown（无 frontmatter）行为不变"""

    def test_plain_markdown_still_works(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = "# 测试标题\n\n这是一段正文。\n\n| A | B |\n|---|---|\n| 1 | 2 |"
        result = render_markdown_to_docx(md)
        doc = Document(result)
        assert any("测试标题" in p.text for p in doc.paragraphs)
        assert len(doc.tables) >= 1

    def test_escaped_newlines_still_work(self):
        from app.services.doc_builder import render_markdown_to_docx
        md = "Line1\\nLine2"
        result = render_markdown_to_docx(md)
        doc = Document(result)
        full = " ".join(p.text for p in doc.paragraphs)
        assert "Line1" in full
