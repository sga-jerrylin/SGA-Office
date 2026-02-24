"""
DOC-X 系列核心业务引擎。
从 main.py 迁移并增强的 Markdown → Docx 渲染器。
支持 YAML frontmatter 驱动的富文档布局（封面、页眉页脚、目录、高亮框、主题色）。
"""

import re
import logging
from io import BytesIO
from typing import Optional, Any

import yaml
import mistune
import requests
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.core.themes import Theme, get_theme

logger = logging.getLogger(__name__)


# =====================================================
#  Frontmatter 解析
# =====================================================

def _parse_frontmatter(markdown_content: str) -> tuple[dict | None, str]:
    """
    从 Markdown 开头提取 YAML frontmatter。
    返回 (config_dict, remaining_body)。无 frontmatter 返回 (None, original)。
    """
    content = markdown_content.strip()
    if not content.startswith("---"):
        return None, markdown_content

    # 找第二个 ---
    end_idx = content.find("---", 3)
    if end_idx == -1:
        return None, markdown_content

    yaml_str = content[3:end_idx].strip()
    body = content[end_idx + 3:].strip()

    try:
        config = yaml.safe_load(yaml_str)
        if not isinstance(config, dict):
            return None, markdown_content
        return config, body
    except yaml.YAMLError:
        return None, markdown_content


# =====================================================
#  页面设置
# =====================================================

def _setup_page(doc: Document) -> None:
    """A4 页面与全局字体设置。"""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = Pt(25)
    style.paragraph_format.first_line_indent = Cm(0.74)


# =====================================================
#  封面页
# =====================================================

def _add_cover_page(doc: Document, cover: dict, theme: Theme) -> None:
    """生成封面页，封面后自动分页。"""
    # 空行撑到页面中部
    for _ in range(6):
        doc.add_paragraph()

    # 主标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(cover["title"])
    title_run.font.size = Pt(26)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(theme.cover_title_color)
    title_run.font.name = '微软雅黑'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 副标题
    subtitle = cover.get("subtitle")
    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_para.paragraph_format.space_before = Pt(12)
        sub_run = sub_para.add_run(subtitle)
        sub_run.font.size = Pt(18)
        sub_run.bold = True
        sub_run.font.name = '微软雅黑'
        sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 空行间隔
    for _ in range(4):
        doc.add_paragraph()

    # 元信息行
    for line in cover.get("meta", []):
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(line)
        meta_run.font.size = Pt(12)
        meta_run.font.color.rgb = RGBColor.from_string(theme.cover_meta_color)
        meta_run.font.name = '微软雅黑'
        meta_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        meta_para.paragraph_format.space_after = Pt(4)
        meta_para.paragraph_format.first_line_indent = Cm(0)

    # 封面后分页
    doc.add_page_break()


# =====================================================
#  页眉页脚
# =====================================================

def _add_header_footer(doc: Document, config: dict) -> None:
    """设置页眉页脚。"""
    section = doc.sections[0]

    # 页眉
    header_text = config.get("header")
    if header_text:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = header_text
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string("808080")

    # 页脚
    footer_type = config.get("footer")
    if footer_type:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if footer_type in ("page_number", "both"):
            # 插入页码域代码
            run = fp.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._element.append(fldChar1)

            run2 = fp.add_run()
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = ' PAGE '
            run2._element.append(instrText)

            run3 = fp.add_run()
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run3._element.append(fldChar2)

        if footer_type in ("custom_text", "both"):
            footer_text_val = config.get("footer_text", "")
            if footer_type == "both":
                fp.add_run("  |  ")
            fp.add_run(footer_text_val)


# =====================================================
#  目录
# =====================================================

def _add_toc(doc: Document) -> None:
    """插入目录占位符（Word 打开时自动刷新）。"""
    toc_title = doc.add_paragraph("目录")
    toc_title.style = doc.styles['Heading 1']

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar2)

    doc.add_page_break()


# =====================================================
#  高亮框（Callout Box）渲染
# =====================================================

_CALLOUT_RE = re.compile(r'^\[!(INFO|NOTE|WARNING)\]\s*(.*)', re.IGNORECASE)


def _render_callout_box(doc: Document, callout_type: str, lines: list[str], theme: Theme) -> None:
    """将高亮框渲染为带底色的单格表格（python-docx 标准着色技巧）。"""
    callout_type_upper = callout_type.upper()
    bg_map = {
        "INFO": theme.callout_info_bg,
        "NOTE": theme.callout_note_bg,
        "WARNING": theme.callout_warning_bg,
    }
    border_map = {
        "INFO": theme.callout_info_border,
        "NOTE": theme.callout_note_border,
        "WARNING": theme.callout_warning_border,
    }
    bg_color = bg_map.get(callout_type_upper, theme.callout_info_bg)
    border_color = border_map.get(callout_type_upper, theme.callout_info_border)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]

    # 设置单元格底色
    tc_pr = cell._element.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), bg_color)
    shading.set(qn('w:val'), 'clear')
    tc_pr.append(shading)

    # 写入内容
    cell._element.clear_content()
    text = "\n".join(lines)
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.first_line_indent = Cm(0)

    # 添加一个空段落作间距
    doc.add_paragraph()


# =====================================================
#  MarkdownToDocx 渲染器 (从 main.py 迁移)
# =====================================================

class MarkdownToDocx:
    """将 mistune 3.x AST 渲染为 python-docx Document"""

    def __init__(self, doc: Document, theme: Theme | None = None):
        self.doc = doc
        self.theme = theme or get_theme()
        self.table = None
        self.row = None
        self.cell = None
        self.list_style = None

    def set_font(self, run, bold=False, italic=False):
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        if bold:
            run.bold = True
        if italic:
            run.italic = True

    def render(self, ast):
        for node in ast:
            self.dispatch(node)

    def dispatch(self, node):
        method_name = f"visit_{node['type']}"
        method = getattr(self, method_name, self.visit_unknown)
        method(node)

    def visit_unknown(self, node):
        if 'children' in node:
            for child in node['children']:
                self.dispatch(child)

    def visit_heading(self, node):
        level = node.get('attrs', {}).get('level', 1)
        text = self.get_text(node)
        heading = self.doc.add_paragraph(style=f'Heading {level}')
        run = heading.add_run(text)
        self.set_font(run, bold=True)
        size_map = {1: 16, 2: 15, 3: 14}
        run.font.size = Pt(size_map.get(level, 12))
        # 主题标题色
        run.font.color.rgb = RGBColor.from_string(self.theme.heading_color)
        heading.paragraph_format.space_before = Pt(22)
        heading.paragraph_format.space_after = Pt(11)

    def visit_paragraph(self, node):
        if len(node.get('children', [])) == 1 and node['children'][0]['type'] == 'image':
            self.visit_image(node['children'][0])
            return
        p = self.doc.add_paragraph()
        if 'children' in node:
            self.render_inline(p, node['children'])
        p.paragraph_format.line_spacing = Pt(25)
        p.paragraph_format.first_line_indent = Cm(0.74)

    def visit_block_code(self, node):
        code = node.get('raw', '')
        p = self.doc.add_paragraph()
        p.style = 'No Spacing'
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(1)

    def visit_list(self, node):
        ordered = node.get('attrs', {}).get('ordered', False)
        self.list_style = 'List Number' if ordered else 'List Bullet'
        for child in node.get('children', []):
            self.visit_list_item(child)
        self.list_style = None

    def visit_list_item(self, node):
        p = self.doc.add_paragraph(style=self.list_style)
        if 'children' in node:
            for child in node['children']:
                if child['type'] in ('paragraph', 'block_text'):
                    self.render_inline(p, child.get('children', []))
                elif child['type'] == 'block_code':
                    run = p.add_run('\n' + child.get('raw', ''))
                    run.font.name = 'Courier New'
                elif child['type'] == 'list':
                    self.visit_list(child)
                elif 'children' in child:
                    self.render_inline(p, child.get('children', []))

    def visit_block_quote(self, node):
        """渲染块引用，检测 [!INFO]/[!NOTE]/[!WARNING] 高亮框。"""
        # 提取块引用的全部文本
        raw_lines = self._extract_blockquote_text(node)
        full_text = "\n".join(raw_lines)

        # 检测 callout 模式
        match = _CALLOUT_RE.match(full_text)
        if match:
            callout_type = match.group(1)
            first_line_rest = match.group(2)
            content_lines = []
            if first_line_rest.strip():
                content_lines.append(first_line_rest.strip())
            # 余下的行
            if len(raw_lines) > 1:
                content_lines.extend(raw_lines[1:])
            _render_callout_box(self.doc, callout_type, content_lines, self.theme)
        else:
            # 普通块引用，渲染为缩进段落
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(full_text)
            self.set_font(run, italic=True)

    def _extract_blockquote_text(self, node) -> list[str]:
        """递归提取 block_quote 节点内所有文本行。"""
        lines = []
        for child in node.get('children', []):
            if child['type'] == 'paragraph':
                # 按 softbreak 拆分行
                lines.extend(self._paragraph_to_lines(child))
            elif child['type'] == 'block_quote':
                lines.extend(self._extract_blockquote_text(child))
            elif 'children' in child:
                text = self.get_text(child)
                lines.extend(text.split('\n'))
            elif 'raw' in child:
                lines.append(child['raw'])
        return lines

    def _paragraph_to_lines(self, para_node) -> list[str]:
        """将段落节点按 softbreak 拆分为多行文本。"""
        lines = []
        current = []
        for child in para_node.get('children', []):
            if child.get('type') == 'softbreak':
                lines.append("".join(current))
                current = []
            else:
                current.append(self.get_text(child))
        if current:
            lines.append("".join(current))
        return lines

    def visit_thematic_break(self, node):
        """水平分割线 (---) 渲染为分页符。"""
        self.doc.add_page_break()

    def visit_table(self, node):
        children = node.get('children', [])
        thead = next((c for c in children if c['type'] == 'table_head'), None)
        tbody = next((c for c in children if c['type'] == 'table_body'), None)

        all_rows = []
        if thead:
            header_cells = thead.get('children', [])
            if header_cells and header_cells[0].get('type') == 'table_cell':
                all_rows.append(header_cells)
            else:
                for row in header_cells:
                    all_rows.append(row.get('children', []))
        if tbody:
            for row in tbody.get('children', []):
                all_rows.append(row.get('children', []))

        if not all_rows:
            return

        col_count = max(len(cells) for cells in all_rows)
        self.table = self.doc.add_table(rows=len(all_rows), cols=col_count)
        self.table.style = "Table Grid"

        for i, cells in enumerate(all_rows):
            self.row = self.table.rows[i]
            for j, cell_node in enumerate(cells):
                self.cell = self.row.cells[j]
                self.cell._element.clear_content()
                p = self.cell.add_paragraph()
                self.render_inline(p, cell_node.get('children', []))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    self.set_font(run)

            # 主题色：表头行着色
            if i == 0:
                for j in range(len(cells)):
                    cell_elem = self.row.cells[j]
                    tc_pr = cell_elem._element.get_or_add_tcPr()
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), self.theme.table_header_bg)
                    shading.set(qn('w:val'), 'clear')
                    tc_pr.append(shading)
                    # 表头字体色
                    for p in cell_elem.paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = RGBColor.from_string(self.theme.table_header_font)
            # 主题色：交替行底色
            elif i % 2 == 0:
                for j in range(len(cells)):
                    cell_elem = self.row.cells[j]
                    tc_pr = cell_elem._element.get_or_add_tcPr()
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), self.theme.table_alt_row_bg)
                    shading.set(qn('w:val'), 'clear')
                    tc_pr.append(shading)

    def visit_image(self, node):
        url = node.get('attrs', {}).get('url', '')
        if not url:
            return

        try:
            response = requests.get(url, timeout=30, headers={
                'User-Agent': 'SGA-Office/1.0'
            })
            if response.status_code != 200:
                self.doc.add_paragraph(f"[图片下载失败: HTTP {response.status_code}]")
                return

            image_stream = BytesIO(response.content)
            img = Image.open(image_stream)
            img_width, img_height = img.size

            # RGBA/P → RGB 转换
            if img.mode in ('RGBA', 'P', 'LA'):
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = rgb_img
            elif img.mode != 'RGB':
                img = img.convert('RGB')

            output_stream = BytesIO()
            img.save(output_stream, format='JPEG', quality=95)
            output_stream.seek(0)

            # 计算合适的宽度
            max_width_cm = 14
            aspect_ratio = img_width / img_height if img_height > 0 else 1
            final_width_cm = max_width_cm
            expected_height_cm = max_width_cm / aspect_ratio
            if expected_height_cm > 18:
                final_width_cm = 18 * aspect_ratio
            final_width_cm = final_width_cm * 0.6

            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(12)

            run = paragraph.add_run()
            inline_shape = run.add_picture(output_stream, width=Cm(final_width_cm))

            # 设置上下型环绕布局
            inline = inline_shape._inline
            anchor = OxmlElement('wp:anchor')
            for attr, val in [
                ('wp:distT', '0'), ('wp:distB', '0'),
                ('wp:distL', '114300'), ('wp:distR', '114300'),
                ('wp:simplePos', '0'), ('wp:relativeHeight', '0'),
                ('wp:behindDoc', '0'), ('wp:locked', '0'),
                ('wp:layoutInCell', '1'), ('wp:allowOverlap', '1'),
            ]:
                anchor.set(qn(attr), val)

            simplePos = OxmlElement('wp:simplePos')
            simplePos.set('x', '0')
            simplePos.set('y', '0')
            anchor.append(simplePos)

            positionH = OxmlElement('wp:positionH')
            positionH.set('relativeFrom', 'column')
            align_elem = OxmlElement('wp:align')
            align_elem.text = 'center'
            positionH.append(align_elem)
            anchor.append(positionH)

            positionV = OxmlElement('wp:positionV')
            positionV.set('relativeFrom', 'paragraph')
            posOffset = OxmlElement('wp:posOffset')
            posOffset.text = '0'
            positionV.append(posOffset)
            anchor.append(positionV)

            for child in inline:
                anchor.append(child)

            wrapTopAndBottom = OxmlElement('wp:wrapTopAndBottom')
            anchor.append(wrapTopAndBottom)

            drawing = inline.getparent()
            drawing.remove(inline)
            drawing.append(anchor)

        except Exception as e:
            logger.error(f"图片处理失败: {e}")
            self.doc.add_paragraph(f"[图片处理失败: {str(e)}]")

    def render_inline(self, paragraph, nodes):
        for node in nodes:
            if node['type'] == 'text':
                run = paragraph.add_run(node.get('raw', ''))
                self.set_font(run)
            elif node['type'] == 'strong':
                run = paragraph.add_run(self.get_text(node))
                self.set_font(run, bold=True)
            elif node['type'] == 'emphasis':
                run = paragraph.add_run(self.get_text(node))
                self.set_font(run, italic=True)
            elif node['type'] == 'codespan':
                run = paragraph.add_run(node.get('raw', ''))
                run.font.name = 'Courier New'
                self.set_font(run)
            elif node['type'] == 'image':
                pass  # inline image skip

    def get_text(self, node):
        if 'raw' in node:
            return node['raw']
        if 'children' in node:
            return "".join([self.get_text(child) for child in node['children']])
        return ""


# =====================================================
#  预处理工具函数
# =====================================================

def _convert_tab_tables_to_markdown(content: str) -> str:
    """将 Tab 分隔的表格转换为 Markdown 表格"""
    lines = content.split('\n')
    result = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if '\t' in line and line.count('\t') >= 1:
            table_lines = []
            while i < len(lines) and '\t' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                for j, tline in enumerate(table_lines):
                    cells = tline.split('\t')
                    md_row = '| ' + ' | '.join(cells) + ' |'
                    result.append(md_row)
                    if j == 0:
                        separator = '| ' + ' | '.join(['---'] * len(cells)) + ' |'
                        result.append(separator)
            else:
                result.extend(table_lines)
        else:
            result.append(line)
            i += 1
    return '\n'.join(result)


def _repair_markdown_table(content: str) -> str:
    """修复 Markdown 表格格式"""
    lines = content.split('\n')
    has_separator = any(
        re.match(r'^\s*\|?[\s:-]{3,}\|[\s:-]{3,}\|?.*$', line)
        for line in lines
    )
    if not has_separator:
        return content

    new_lines = []
    for line in lines:
        stripped = line.strip()
        if '|' in stripped and len(stripped) > 1:
            if not stripped.startswith('|'):
                stripped = '| ' + stripped
            if not stripped.endswith('|'):
                stripped = stripped + ' |'
            new_lines.append(stripped)
        else:
            new_lines.append(line)
    return '\n'.join(new_lines)


# =====================================================
#  DOC-01: render_markdown_to_docx
# =====================================================

def render_markdown_to_docx(
    markdown_content: str,
) -> BytesIO:
    """
    将 Markdown 文本渲染为标准版式 Word 文档。
    支持 YAML frontmatter 驱动的封面、页眉页脚、目录、高亮框和主题色。
    无 frontmatter 时行为与原版完全一致。

    Args:
        markdown_content: Markdown 原文（可含 YAML frontmatter）

    Returns:
        BytesIO 对象，包含生成的 .docx 数据
    """
    # 1. 解析 frontmatter
    config, body = _parse_frontmatter(markdown_content)

    # 2. 加载主题
    theme_name = config.get("theme") if config else None
    theme = get_theme(theme_name)

    # 3. 预处理正文
    content = body
    if content.startswith('\\#'):
        content = content[1:]
    content = content.replace('\\n', '\n')
    content = _convert_tab_tables_to_markdown(content)
    content = _repair_markdown_table(content)

    # 4. 创建文档
    doc = Document()
    _setup_page(doc)

    # 5. 封面页
    if config and config.get("cover"):
        cover = config["cover"]
        if not cover.get("title"):
            from app.core.error_hints import build_agent_hint, ErrorType
            raise ValueError(str(build_agent_hint(
                ErrorType.MISSING_FIELD, field="cover.title",
                message="封面配置中 title 是必填字段"
            )))
        _add_cover_page(doc, cover, theme)

    # 6. 目录
    if config and config.get("toc"):
        _add_toc(doc)

    # 7. 页眉页脚
    if config:
        _add_header_footer(doc, config)

    # 8. 解析 & 渲染正文
    markdown = mistune.create_markdown(renderer=None, plugins=['table'])
    ast = markdown(content)
    renderer = MarkdownToDocx(doc, theme=theme)
    renderer.render(ast)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# =====================================================
#  DOC-02: fill_docx_template
# =====================================================

def fill_docx_template(
    template_bytes: bytes,
    variables: dict[str, Any],
) -> BytesIO:
    """
    将 Word 模板中的 {{ 占位符 }} 替换为实际变量值。

    Args:
        template_bytes: 模板 .docx 文件的原始字节
        variables: key-value 变量字典，key 对应模板中的占位符名称

    Returns:
        BytesIO 对象，包含替换后的 .docx 数据
    """
    doc = Document(BytesIO(template_bytes))

    for paragraph in doc.paragraphs:
        _replace_placeholders_in_paragraph(paragraph, variables)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_placeholders_in_paragraph(paragraph, variables)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _replace_placeholders_in_paragraph(
    paragraph,
    variables: dict[str, Any],
) -> None:
    """在单个段落中替换 {{ key }} 占位符。"""
    for key, value in variables.items():
        placeholder = "{{" + key + "}}"
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))
