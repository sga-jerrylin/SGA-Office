from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import os
import re
from urllib.parse import quote
import json
import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment, Font, Border, Side, PatternFill
from io import BytesIO
from qcloud_cos import CosConfig, CosS3Client
import mistune
import requests

app = Flask(__name__)

# 腾讯云COS配置（从环境变量读取）
SECRET_ID = os.environ.get('COS_SECRET_ID', '')
SECRET_KEY = os.environ.get('COS_SECRET_KEY', '')
REGION = os.environ.get('COS_REGION', 'ap-guangzhou')
BUCKET_NAME = os.environ.get('COS_BUCKET_NAME', '')

config = CosConfig(Region=REGION, SecretId=SECRET_ID, SecretKey=SECRET_KEY)
cos_client = CosS3Client(config)

####################################
#          Markdown 渲染器
####################################
class MarkdownToDocx:
    def __init__(self, doc):
        self.doc = doc
        self.table = None
        self.row = None
        self.cell = None
        self.list_style = None

    def set_font(self, run, bold=False, italic=False):
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        if bold:
            run.bold = True
        if italic:
            run.italic = True

    def render(self, ast):
        for node in ast:
            self.dispatch(node)

    def dispatch(self, node):
        method_name = f"visit_{node['type']}"
        method = getattr(self, method_name, self.visit_unknown)
        method(node)

    def visit_unknown(self, node):
        if 'children' in node:
            for child in node['children']:
                self.dispatch(child)

    def visit_heading(self, node):
        level = node.get('attrs', {}).get('level', 1)
        text = self.get_text(node)
        heading = self.doc.add_paragraph(style=f'Heading {level}')
        run = heading.add_run(text)
        self.set_font(run, bold=True)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(15)
        elif level == 3:
            run.font.size = Pt(14)
        heading.paragraph_format.space_before = Pt(22)
        heading.paragraph_format.space_after = Pt(11)

    def visit_paragraph(self, node):
        # 如果段落里只有图片，特殊处理
        if len(node.get('children', [])) == 1 and node['children'][0]['type'] == 'image':
            self.visit_image(node['children'][0])
            return

        p = self.doc.add_paragraph()
        if 'children' in node:
            self.render_inline(p, node['children'])
        p.paragraph_format.line_spacing = Pt(25)
        p.paragraph_format.first_line_indent = Cm(0.74)

    def visit_block_code(self, node):
        code = node.get('raw', '')
        p = self.doc.add_paragraph()
        p.style = 'No Spacing' 
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(1)

    def visit_list(self, node):
        ordered = node.get('attrs', {}).get('ordered', False)
        self.list_style = 'List Number' if ordered else 'List Bullet'
        for child in node.get('children', []):
            self.visit_list_item(child)
        self.list_style = None

    def visit_list_item(self, node):
        p = self.doc.add_paragraph(style=self.list_style)
        if 'children' in node:
             for child in node['children']:
                 if child['type'] == 'paragraph':
                     self.render_inline(p, child.get('children', []))
                 elif child['type'] == 'block_text':
                     # mistune 3.x 使用 block_text 而不是 paragraph
                     self.render_inline(p, child.get('children', []))
                 elif child['type'] == 'block_code':
                     run = p.add_run('\n' + child.get('raw', ''))
                     run.font.name = 'Courier New'
                 elif child['type'] == 'list':
                     # 处理嵌套列表
                     self.visit_list(child)
                 else:
                     # 尝试处理其他包含 children 的节点
                     if 'children' in child:
                         self.render_inline(p, child.get('children', []))

    def visit_table(self, node):
        print(f"[DEBUG] visit_table called")
        children = node.get('children', [])
        thead = next((c for c in children if c['type'] == 'table_head'), None)
        tbody = next((c for c in children if c['type'] == 'table_body'), None)

        print(f"[DEBUG] thead found: {thead is not None}, tbody found: {tbody is not None}")

        # 构建统一的行结构
        # thead 的 children 直接是 table_cell 列表
        # tbody 的 children 是 table_row 列表，每个 table_row 包含 table_cell
        all_rows = []  # 每个元素是 cell 列表

        if thead:
            # thead.children 直接是 table_cell 列表
            header_cells = thead.get('children', [])
            print(f"[DEBUG] thead children count: {len(header_cells)}")
            if header_cells and header_cells[0].get('type') == 'table_cell':
                all_rows.append(header_cells)
            else:
                # 兼容旧格式：thead.children 是 table_row 列表
                for row in header_cells:
                    all_rows.append(row.get('children', []))

        if tbody:
            # tbody.children 是 table_row 列表
            tbody_children = tbody.get('children', [])
            print(f"[DEBUG] tbody children count: {len(tbody_children)}")
            for row in tbody_children:
                all_rows.append(row.get('children', []))

        print(f"[DEBUG] all_rows count: {len(all_rows)}")

        if not all_rows:
            print(f"[DEBUG] No rows found, returning")
            return

        # Calculate max columns to ensure all data fits
        col_count = max(len(cells) for cells in all_rows)
        print(f"[DEBUG] Creating table with {len(all_rows)} rows and {col_count} cols")
        self.table = self.doc.add_table(rows=len(all_rows), cols=col_count)
        self.table.style = "Table Grid"

        for i, cells in enumerate(all_rows):
            self.row = self.table.rows[i]
            for j, cell_node in enumerate(cells):
                self.cell = self.row.cells[j]
                self.cell._element.clear_content()
                p = self.cell.add_paragraph()
                self.render_inline(p, cell_node.get('children', []))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    self.set_font(run)

        print(f"[DEBUG] Table created successfully")

    def visit_image(self, node):
        url = node.get('attrs', {}).get('url', '')
        if not url:
            print(f"[DEBUG] Image URL is empty")
            return

        print(f"[DEBUG] Downloading image: {url}")

        try:
            # 下载图片
            response = requests.get(url, timeout=30, headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            })

            if response.status_code != 200:
                print(f"[DEBUG] Image download failed with status: {response.status_code}")
                self.doc.add_paragraph(f"[图片下载失败: HTTP {response.status_code}]")
                return

            image_data = response.content
            print(f"[DEBUG] Image downloaded, size: {len(image_data)} bytes")

            # 验证并处理图片
            from PIL import Image
            image_stream = BytesIO(image_data)
            try:
                img = Image.open(image_stream)
                img_width, img_height = img.size
                img_format = img.format
                img_mode = img.mode
                print(f"[DEBUG] Image info: {img_width}x{img_height}, format: {img_format}, mode: {img_mode}")

                # 关键修复：将 RGBA/P 等模式转换为 RGB，避免 Word 显示问题
                if img_mode in ('RGBA', 'P', 'LA'):
                    print(f"[DEBUG] Converting {img_mode} to RGB")
                    # 创建白色背景
                    rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                    if img_mode == 'P':
                        img = img.convert('RGBA')
                    rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = rgb_img
                elif img_mode != 'RGB':
                    print(f"[DEBUG] Converting {img_mode} to RGB")
                    img = img.convert('RGB')

                # 保存为 JPEG 格式（兼容性更好）
                output_stream = BytesIO()
                img.save(output_stream, format='JPEG', quality=95)
                output_stream.seek(0)
                print(f"[DEBUG] Image converted to JPEG, size: {len(output_stream.getvalue())} bytes")

            except Exception as img_err:
                print(f"[DEBUG] Invalid image format: {img_err}")
                self.doc.add_paragraph(f"[图片格式无效: {url}]")
                return

            # 计算合适的宽度（固定使用 14cm，适合 A4 页面）
            max_width_cm = 14
            aspect_ratio = img_width / img_height if img_height > 0 else 1

            # 简化逻辑：所有图片统一缩放到最大14cm宽
            final_width_cm = max_width_cm

            # 如果图片太高（高度会超过18cm），则按高度限制
            expected_height_cm = max_width_cm / aspect_ratio
            if expected_height_cm > 18:
                final_width_cm = 18 * aspect_ratio

            # 再缩小到 60%，效果更好
            final_width_cm = final_width_cm * 0.6
            final_width = Cm(final_width_cm)

            print(f"[DEBUG] Final width: {final_width}")

            # 创建新段落并添加图片
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(12)

            run = paragraph.add_run()
            inline_shape = run.add_picture(output_stream, width=final_width)

            # 设置图片为"上下型环绕"布局
            from docx.oxml.ns import nsmap, qn
            from docx.oxml import OxmlElement

            # 获取 inline 元素并转换为 anchor（浮动）
            inline = inline_shape._inline

            # 创建 anchor 元素替代 inline
            anchor = OxmlElement('wp:anchor')
            anchor.set(qn('wp:distT'), '0')
            anchor.set(qn('wp:distB'), '0')
            anchor.set(qn('wp:distL'), '114300')
            anchor.set(qn('wp:distR'), '114300')
            anchor.set(qn('wp:simplePos'), '0')
            anchor.set(qn('wp:relativeHeight'), '0')
            anchor.set(qn('wp:behindDoc'), '0')
            anchor.set(qn('wp:locked'), '0')
            anchor.set(qn('wp:layoutInCell'), '1')
            anchor.set(qn('wp:allowOverlap'), '1')

            # simplePos
            simplePos = OxmlElement('wp:simplePos')
            simplePos.set('x', '0')
            simplePos.set('y', '0')
            anchor.append(simplePos)

            # positionH - 水平居中
            positionH = OxmlElement('wp:positionH')
            positionH.set('relativeFrom', 'column')
            align = OxmlElement('wp:align')
            align.text = 'center'
            positionH.append(align)
            anchor.append(positionH)

            # positionV - 垂直相对于段落
            positionV = OxmlElement('wp:positionV')
            positionV.set('relativeFrom', 'paragraph')
            posOffset = OxmlElement('wp:posOffset')
            posOffset.text = '0'
            positionV.append(posOffset)
            anchor.append(positionV)

            # 复制 inline 的子元素
            for child in inline:
                anchor.append(child)

            # 添加 wrapTopAndBottom（上下型环绕）
            wrapTopAndBottom = OxmlElement('wp:wrapTopAndBottom')
            anchor.append(wrapTopAndBottom)

            # 替换 inline 为 anchor
            drawing = inline.getparent()
            drawing.remove(inline)
            drawing.append(anchor)

            print(f"[DEBUG] Image added with Top-Bottom wrap layout")

        except requests.exceptions.Timeout:
            print(f"[DEBUG] Image download timeout: {url}")
            self.doc.add_paragraph(f"[图片下载超时: {url}]")
        except requests.exceptions.RequestException as e:
            print(f"[DEBUG] Image download error: {e}")
            self.doc.add_paragraph(f"[图片下载失败: {url}]")
        except Exception as e:
            print(f"[DEBUG] Unexpected error adding image: {e}")
            import traceback
            traceback.print_exc()
            self.doc.add_paragraph(f"[图片处理失败: {str(e)}]")

    def render_inline(self, paragraph, nodes):
        for node in nodes:
            if node['type'] == 'text':
                run = paragraph.add_run(node.get('raw', ''))
                self.set_font(run)
            elif node['type'] == 'strong':
                run = paragraph.add_run(self.get_text(node))
                self.set_font(run, bold=True)
            elif node['type'] == 'emphasis':
                run = paragraph.add_run(self.get_text(node))
                self.set_font(run, italic=True)
            elif node['type'] == 'image':
                 # Inline image - simplified
                 pass
            elif node['type'] == 'codespan':
                 run = paragraph.add_run(node.get('raw', ''))
                 run.font.name = 'Courier New'
                 self.set_font(run)

    def get_text(self, node):
        if 'raw' in node:
            return node['raw']
        if 'children' in node:
            return "".join([self.get_text(child) for child in node['children']])
        return ""

#         Docx 生成接口
####################################
@app.route('/')
def home():
    return "Flask 服务器运行成功！"

@app.route('/generate-doc', methods=['POST'])
def generate_doc():
    data = request.json

    # 调试日志：打印收到的原始数据
    print(f"[DEBUG] Received data: {data}")
    print(f"[DEBUG] Data type: {type(data)}")

    if not data:
        return jsonify({'error': '没有提供数据'}), 400

    # 确保 filename 和 content 是字符串类型
    filename_input = str(data.get("filename", "默认文档")).strip()
    content = str(data.get("content", ""))

    print(f"[DEBUG] filename_input: {filename_input}")
    print(f"[DEBUG] content (first 100 chars): {content[:100] if content else 'EMPTY'}")

    # 处理可能的转义字符和特殊前缀
    # 有些系统可能会对 # 开头的内容进行转义
    if content.startswith('\\#'):
        content = content[1:]  # 去掉转义符

    # 处理 \n 字面量转换为真正的换行符
    content = content.replace('\\n', '\n')

    if not content or content.strip() == '' or content == 'None':
         print(f"[DEBUG] Content is empty or None!")
         return jsonify({'error': '内容不能为空'}), 400

    try:
        doc = Document()

        # ========== 页面设置 ==========
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

        # ========== 全局字体设置 ==========
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = Pt(25)
        style.paragraph_format.first_line_indent = Cm(0.74)

        # ========== Markdown 解析与渲染 ==========
        markdown = mistune.create_markdown(renderer=None, plugins=['table'])
        ast = markdown(content)
        
        renderer = MarkdownToDocx(doc)
        renderer.render(ast)

        # ========== 生成文件名 ==========
        clean_title = re.sub(r'[\\/:*?"<>|\s]', '', filename_input)[:30]
        if not clean_title:
            clean_title = "无标题文档"
        current_date = datetime.now().strftime("%Y%m%d")
        filename = f"{clean_title}_{current_date}.docx"

        # ========== 保存并上传到COS ==========
        temp_file = f"temp_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        doc.save(temp_file)
        
        cos_path = f"documents/{filename}"
        try:
            cos_client.upload_file(
                Bucket=BUCKET_NAME,
                LocalFilePath=temp_file,
                Key=cos_path
            )
        except Exception as upload_error:
            app.logger.error(f"上传失败: {str(upload_error)}")
            return jsonify({'error': '文件上传失败'}), 500
        finally:
            if os.path.exists(temp_file):
                os.remove(temp_file)
        
        file_url = f"https://{BUCKET_NAME}.cos.{REGION}.myqcloud.com/{quote(cos_path)}"
        return jsonify({"message": "生成成功", "file_url": file_url})
    
    except Exception as e:
        app.logger.error(f"生成文档错误: {str(e)}")
        return jsonify({'error': str(e)}), 500

####################################
#         Excel 生成接口
####################################
def validate_json_data(json_data):
    """验证JSON数据结构"""
    required_fields = ['title', 'data']
    if not all(field in json_data for field in required_fields):
        raise ValueError("Missing required fields in JSON data")
    
    if not json_data['data'] or not isinstance(json_data['data'], list):
        raise ValueError("Data must be a non-empty array")
    
    if not json_data['data'][0] or not isinstance(json_data['data'][0], list):
        raise ValueError("Headers must be a non-empty array")

def calculate_title_style(title):
    """根据标题长度计算字体大小和行高"""
    title_length = len(title)
    if title_length <= 15:
        return {'size': 16, 'height': 25}
    elif 15 < title_length <= 25:
        return {'size': 14, 'height': 35}
    elif 25 < title_length <= 40:
        return {'size': 12, 'height': 45}
    else:
        return {'size': 11, 'height': 55}

def apply_cell_style(cell, is_header=False, is_summary=False):
    """应用单元格样式"""
    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    cell.border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    if is_header:
        cell.font = Font(bold=True, size=11)
        cell.fill = PatternFill(start_color='F2F2F2', end_color='F2F2F2', fill_type='solid')
    if is_summary:
        cell.font = Font(bold=True, size=11)
        cell.fill = PatternFill(start_color='E6F3FF', end_color='E6F3FF', fill_type='solid')

def generate_excel(json_data):
    """生成Excel文件"""
    validate_json_data(json_data)
    
    wb = openpyxl.Workbook()
    sheet = wb.active
    title = json_data['title']
    title_style = calculate_title_style(title)
    sheet['A1'] = title
    title_cell = sheet['A1']
    title_cell.font = Font(size=title_style['size'], bold=True)
    title_cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    data = json_data['data']
    headers = data[0]
    data_rows = data[1:]
    end_col = get_column_letter(len(headers))
    sheet.merge_cells(f'A1:{end_col}1')
    sheet.row_dimensions[1].height = title_style['height']
    for col_num, header in enumerate(headers, 1):
        cell = sheet.cell(row=2, column=col_num, value=header)
        apply_cell_style(cell, is_header=True)
    for row_num, row_data in enumerate(data_rows, 3):
        is_summary = row_data[0] in ['合计', '总计', 'Total']
        for col_num, value in enumerate(row_data, 1):
            cell = sheet.cell(row=row_num, column=col_num, value=value)
            apply_cell_style(cell, is_summary=is_summary)
    for col_num in range(1, len(headers) + 1):
        max_len = 0
        col_letter = get_column_letter(col_num)
        for cell in sheet[col_letter]:
            try:
                if cell.value:
                    cell_len = len(str(cell.value))
                    max_len = max(max_len, cell_len)
            except:
                pass
        adjusted_width = (max_len + 2) * 1.2
        sheet.column_dimensions[col_letter].width = adjusted_width
    if 'metadata' in json_data:
        metadata = json_data['metadata']
        metadata_row = sheet.max_row + 2
        if 'summary' in metadata:
            sheet.cell(row=metadata_row, column=1, value=metadata['summary'])
            sheet.merge_cells(f'A{metadata_row}:{end_col}{metadata_row}')
        if 'timestamp' in metadata:
            timestamp_row = metadata_row + 1
            try:
                timestamp_date = datetime.strptime(metadata['timestamp'], "%Y-%m-%dT%H:%M:%SZ").strftime("%Y-%m-%d")
            except:
                timestamp_date = datetime.now().strftime("%Y-%m-%d")
            sheet.cell(row=timestamp_row, column=1, value=f"生成时间：{timestamp_date}")
            sheet.merge_cells(f'A{timestamp_row}:{end_col}{timestamp_row}')
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output

@app.route('/generate_excel', methods=['POST'])
def generate_excel_route():
    try:
        json_data = request.get_json()
        if not json_data:
            return jsonify({"error": "No JSON data received"}), 400
        
        filename_input = json_data.get("filename", "未命名表格").strip()
        content = json_data.get("content", "")

        if not content:
             return jsonify({'error': '内容不能为空'}), 400

        # Parse Markdown to find table
        markdown = mistune.create_markdown(renderer=None, plugins=['table'])
        ast = markdown(content)
        
        table_node = None
        # Find the first table node
        # Mistune 3 AST is a list of nodes
        for node in ast:
            if node['type'] == 'table':
                table_node = node
                break
        
        if not table_node:
            return jsonify({'error': '未找到表格内容'}), 400

        # Extract data from table node
        headers = []
        rows = []
        
        children = table_node.get('children', [])
        thead = next((c for c in children if c['type'] == 'table_head'), None)
        tbody = next((c for c in children if c['type'] == 'table_body'), None)

        def extract_text(node):
            if 'raw' in node:
                return node['raw']
            if 'children' in node:
                return "".join([extract_text(child) for child in node['children']])
            return ""

        if thead:
            # Assuming single row in thead for now
            for row in thead.get('children', []):
                for cell in row.get('children', []):
                    headers.append(extract_text(cell))
        
        if tbody:
            for row in tbody.get('children', []):
                current_row = []
                for cell in row.get('children', []):
                    current_row.append(extract_text(cell))
                rows.append(current_row)

        if not headers and not rows:
             return jsonify({'error': '表格为空'}), 400

        # Construct data for generate_excel
        # Ensure headers is a list, even if empty (though valid markdown table usually has headers)
        excel_input_data = {
            'title': filename_input,
            'data': [headers] + rows
        }

        print(f"Received request at {datetime.now()}")
        # print("Parsed Excel Data:", json.dumps(excel_input_data, ensure_ascii=False, indent=2))

        excel_data = generate_excel(excel_input_data)
        
        clean_title = re.sub(r'[\\/:*?"<>|\s]', '', filename_input)[:30]
        if not clean_title:
            clean_title = "未命名表格"
        current_date = datetime.now().strftime("%Y%m%d")
        filename = f"{clean_title}_{current_date}.xlsx"
        temp_file = f"temp_{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
        
        with open(temp_file, 'wb') as f:
            f.write(excel_data.getvalue())
            
        try:
            cos_path = f"excel_documents/{filename}"
            cos_client.upload_file(
                Bucket=BUCKET_NAME,
                LocalFilePath=temp_file,
                Key=cos_path
            )
            file_url = f"https://{BUCKET_NAME}.cos.{REGION}.myqcloud.com/{quote(cos_path)}"
            return jsonify({
                "message": "Excel文件生成成功",
                "file_url": file_url,
                "filename": filename
            })
        except Exception as upload_error:
            app.logger.error(f"上传失败: {str(upload_error)}")
            return jsonify({'error': '文件上传失败'}), 500
        finally:
            if os.path.exists(temp_file):
                os.remove(temp_file)
    except ValueError as ve:
        return jsonify({"error": str(ve)}), 400
    except Exception as e:
        print(f"Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": "Internal server error"}), 500

@app.route('/health', methods=['GET'])
def health_check():
    """健康检查接口"""
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "service": "Excel & Docx Generator"
    })

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5101, debug=True)
